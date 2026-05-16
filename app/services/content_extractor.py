import base64
import io

import fitz  # PyMuPDF
import logfire
from docx import Document

from app.utils.text_utils import contains_resume_keywords


class ContentExtractor:
    @staticmethod
    async def extract_text_from_pdf(pdf_bytes: bytes) -> str:
        """Extract text from PDF using layout-aware block reading."""
        with logfire.span("extract_text_from_pdf", size=len(pdf_bytes)):
            try:
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                text_content = []
                for page in doc:
                    blocks = page.get_text("blocks")
                    blocks.sort(key=lambda b: (b[1], b[0]))
                    for b in blocks:
                        text_content.append(b[4])
                doc.close()
                logfire.info(f"Extracted {len(text_content)} text blocks from PDF")
                return "\n".join(text_content)
            except Exception as e:
                logfire.error(f"PDF text extraction failed: {e}", error=str(e))
                return ""

    @staticmethod
    async def extract_text_from_docx(docx_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        with logfire.span("extract_text_from_docx", size=len(docx_bytes)):
            try:
                doc = Document(io.BytesIO(docx_bytes))
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                            
                logfire.info(f"Extracted {len(full_text)} paragraphs/cells from DOCX")
                return "\n".join(full_text)
            except Exception as e:
                logfire.error(f"DOCX extraction failed: {e}", error=str(e))
                return ""

    @staticmethod
    def render_pdf_to_images(pdf_bytes: bytes) -> list[str]:
        """Render PDF pages to base64 encoded images for OCR."""
        with logfire.span("render_pdf_to_images", size=len(pdf_bytes)):
            images_base64 = []
            try:
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                for _i, page in enumerate(doc):
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    encoded = base64.b64encode(img_data).decode("utf-8")
                    images_base64.append(encoded)
                doc.close()
                logfire.info(f"Rendered {len(images_base64)} pages to images")
                return images_base64
            except Exception as e:
                logfire.error(f"Failed to render PDF to images: {e}", error=str(e))
                return []

    @staticmethod
    def prepare_image_for_ocr(image_bytes: bytes) -> str:
        """Convert image bytes to base64."""
        return base64.b64encode(image_bytes).decode("utf-8")

    @staticmethod
    async def extract_text(file_bytes: bytes, filename: str) -> str:
        """Extract text from various file types."""
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            return await ContentExtractor.extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            return await ContentExtractor.extract_text_from_docx(file_bytes)
        elif filename_lower.endswith(".txt"):
            try:
                return file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1")
        return ""

    @classmethod
    def validate_content_quality(cls, text: str) -> bool:
        """Check if the extracted text is sufficient and looks like a resume."""
        if not text or len(text.strip()) < 100:
            logfire.debug("Text too short for quality check", text_length=len(text) if text else 0)
            return False
        
        # Check for anchor keywords
        is_valid = contains_resume_keywords(text)
        if not is_valid:
            logfire.debug("Resume keywords not found in text")
        return is_valid
